from tkinter import *
from PIL import ImageTk, Image
from tkinter import filedialog
from tkinter.filedialog import askopenfile
import mysql.connector
from requests import options
from PIL import Image
from docx import Document
from docx.opc.coreprops import CoreProperties
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from datetime import date
from docx import Document
import tkinter as tk
from tkinter import filedialog
from docx import Document
from tkPDFViewer import tkPDFViewer as pdf
import os
import docx
from tkinter import scrolledtext
from docx2pdf import convert
import shutil



filename_path = None
filename_path_1 = None

def fetching():

    mydb = mysql.connector.connect(host="localhost",username="root",password="",db="vehicle_detection_system")
    mycursor = mydb.cursor()
    sql = f"select first_name, last_name, phone, course, university, vehicle, license from vehicle_registration_data where vehicle = '{name_entry.get()}'"
    mycursor.execute(sql)

    records = mycursor.fetchone()

    first_name =  records[0]
    last_name = records[1]
    phone =  records[2]
    course = records[3]
    university_name = records[4]
    vehicle_no = records[5]
    License_no = records[6]
    current_date = date.today()
    mydb.commit()

    logo = 'college_logo.jpg'
    doc = Document()    

    # logo image placement
    logo = doc.add_picture(logo, width=Inches(2.52), height=Inches(1.5))
    logo_placement = doc.paragraphs[-1] 
    logo_placement.alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_paragraph = doc.add_paragraph()
    
    new_paragraph.add_run(f'''
    
                                 "Application for Vehicle Approval for Campus Parking"

                            

Date: {current_date}

    

Name of the Student: {first_name} {last_name}

Contact no: {phone}     

Institute Name: {course},{university_name}

Vehicle number: {vehicle_no}	   

License no: {License_no}

Is the vehicle insured? Yes / No

    
P.S.: kindly note parking is at your own risk; management is not liable for any loss/damage/theft.




			 
Signature of Staff member		         Signature of the Head of the Institute

    
    ''')


    doc.save(f"Form/{first_name} {last_name} {vehicle_no}.docx")
    global filename_path
    global filename_path_1
    filename_path = f"Form/{first_name} {last_name} {vehicle_no}.docx"
    filename_path_1 = f"Form/{first_name} {last_name} {vehicle_no}.pdf"

    convert(filename_path,f"Form/{first_name} {last_name} {vehicle_no}.pdf")


def open():
    global filename_path_1
    v1 = pdf.ShowPdf()
    v2 = v1.pdf_view(root, pdf_location = filename_path_1, width = 70, height = 38)
    v2.place(x=700, y = 38)


def call_both_functions():
    fetching()
    open()


def back():
    root.destroy()
    import Vehicle_Register_Form

def main_page():
    root.destroy()
    import main_landing_page

def download():
    pdf_file_path = filename_path_1

    root = Tk()
    root.withdraw()
    destination_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
    root.destroy()

    # Copy the PDF file to the selected destination path
    if destination_path:
        shutil.copyfile(pdf_file_path, destination_path)
        print("PDF saved successfully!")


root = Tk()
root.title("Preview Form")
root.resizable(False, False)
root.geometry("1400x700+60+20")

img2 =Image.open("approval_bg.jpg")
bg2 = ImageTk.PhotoImage(img2)

# Add image
label2 = Label(root, image=bg2, borderwidth=0)
label2.place(x = 0,y = 0)

canvas = Canvas(root, width=1200, height=650)
canvas.pack()
png = PhotoImage(file = "white_box .png") # Just an example
canvas.create_image(0, 0, image = png, anchor = "nw")
canvas.place(relx=0.5, rely=0.5, anchor=CENTER)
canvas.create_line(525, 200, 525, 510, fill="black", width=3)


enrty_name_1 = Image.open("input_image.png")
resize = enrty_name_1.resize((250,50))
new_pic = ImageTk.PhotoImage(resize)
name_label = Label(root, text="Enter the Vehicle Number :- ", bg="white", font=('bold',14))
name_label.place(x=240,y=70)
enrty_label = Label(root, image=new_pic, border=0)
enrty_label.place(x=250, y=120)
name_entry = Entry(root,width=15,border=0,font=('bold',18))
name_entry.place(x=265,y=130)


upload = PhotoImage(file="preview_logo.png")
img_label_4 = Label(image=upload)
b1 = Button(root, image = upload, text='Upload File',borderwidth=0,command = lambda:call_both_functions())
b1.place(x = 250, y = 200)


upload2 = PhotoImage(file="download_icon.png")
img_label_4 = Label(image=upload2)
b1 = Button(root, image = upload2, text='Upload File',borderwidth=0, command=lambda:download())
b1.place(x = 300, y = 300)

upload3 = PhotoImage(file="back_logo.png")
img_label_5 = Label(image=upload3)
b1 = Button(root, image = upload3, text='Back',borderwidth=0, command=lambda:back())
b1.place(x = 105, y = 580)

upload4 = PhotoImage(file="main_page_logo.png")
img_label_5 = Label(image=upload4)
b1 = Button(root, image = upload4, text='Back',borderwidth=0, command=lambda:main_page())
b1.place(x = 370, y = 580)


root.mainloop()